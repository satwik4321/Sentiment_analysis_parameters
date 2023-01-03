#!/usr/bin/env python
# coding: utf-8

# In[1]:


import pandas as pd
from bs4 import BeautifulSoup
import requests
import docx
from docx import Document
import os
import re
import openpyxl
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter


# In[2]:


df=pd.read_excel(r"C:\Users\sathw\Downloads\Input (2).xlsx")


# In[3]:


n=int(df['URL_ID'][0])
s=n-1
l=[]
for i in df['URL']:
    r = requests.get(str(i), headers={"User-Agent": "XY"})
    if r.status_code==404:
        n=n+1
        continue
    soup=BeautifulSoup(r.content,"html.parser")
    title=soup.find("h1").get_text()
    text=soup.find_all("p")
    doc=docx.Document()
    doc.add_paragraph(title)
    for para in text:
        text1=para.get_text()
        doc.add_paragraph(text1)
    doc.save(r"C:\Users\sathw\Downloads\Scrapped_text\{}.docx".format(n))
    n=n+1


# In[4]:


sw=[]
md=[]
pathsw=r"C:\Users\sathw\Downloads\Stopwords"
pathmd=r"C:\Users\sathw\Downloads\MasterDictionary"
files=os.listdir(pathsw)
for docs in files:
    d=os.path.join(pathsw,docs)
    sw.append(str(d))
files=os.listdir(pathmd)
for docs in files:
    d=os.path.join(pathmd,docs)
    md.append(str(d))
filenames=os.listdir(r'C:\Users\sathw\Downloads\Scrapped_text')
stopwords=[]
for file in sw:
    f=open(file)
    for word in f:
        word=word.lower()
        l=word.split(' ')
        word=l[0]
        word=word.strip('\n')
        word=word.strip(',')
        stopwords.append(word)
print(stopwords)


# In[5]:


def syllable(word):
    vowels=['a','e','i','o','u']
    temp='a'
    cnt=0
    for i in word:
        if i in vowels:
            if temp not in vowels:
                cnt+=1
        temp=word
    if cnt>2:
        return True
    return False


# In[6]:


def get_metrics(doc,ind):
    string=""
    for para in doc.paragraphs:
        string=string+(para.text.lower())+"\n"
    sent=1
    for i in string:
        if i=='\n':
            sent+=1
    string=re.sub(r'[.,?!;]','',string)
    words=string.split()
    cnt=0
    n_char=0
    complex_words=0
    for word in words:
        cnt+=1
        l=word.strip(',')
        l=l.strip('.')
        l=l.strip('!')
        l=l.strip('?')
        l=l.strip()
        n_char+=len(l)
        if(syllable(word)):
            complex_words+=1
        if l in stopwords:
            words.remove(word)
    positive_score=0
    negative_score=0
    f=open(md[0])
    for word in f:
        word=word.strip()
        word=word.strip('\n')
        if word in words:
            positive_score+=1
    f=open(md[1])
    for word in f:
        word=word.strip()
        word=word.strip('\n')
        if word in words:
            negative_score+=1
    polarity_score=(positive_score-negative_score)/((positive_score+negative_score)+0.000001)
    subjectivity_score=(positive_score+negative_score)/(len(words)+0.000001)
    pronoun = re.compile(r'I|we|my|ours|us',re.I)
    pronouns = pronoun.findall(string)
    avg_sent_len=(cnt/sent)
    percent_comp_words=(complex_words/cnt)*100
    fog_index=0.4*((complex_words/cnt)+(cnt/sent))
    syllable_per_word=(complex_words/cnt)
    personal_pronouns=len(pronouns)
    avg_word_len=n_char/cnt
    parameters=[positive_score,negative_score,polarity_score,subjectivity_score,avg_sent_len,percent_comp_words,fog_index,avg_sent_len,complex_words,cnt,syllable_per_word,personal_pronouns,avg_word_len]
    for j,value in enumerate(parameters):
        char=get_column_letter(j+3)
        ws[char+str(ind+1)]=str(value)
    wb.save(r"C:\Users\sathw\Downloads\Input (2).xlsx")


# In[7]:


wb=load_workbook(r"C:\Users\sathw\Downloads\Input (2).xlsx")
ws=wb.active
columns=['POSITIVE SCORE','NEGATIVE SCORE','POLARITY SCORE','SUBJECTIVITY SCORE','AVG SENTENCE LENGTH','PERCENTAGE OF COMPLEX WORDS','FOG INDEX','AVG NUMBER OF WORDS PER SENTENCE','COMPLEX WORD COUNT','WORD COUNT','SYLLABLE PER WORD','PERSONAL PRONOUNS','AVG WORD LENGTH']
for i,s in enumerate(columns):
    char=get_column_letter(i+3)
    ws[char+str(1)].value=str(s)
wb.save(r"C:\Users\sathw\Downloads\Input (2).xlsx")
j=0
for filename in filenames:
    doc=Document(r"C:\\Users\\sathw\\Downloads\\Scrapped_text\\"+str(filename))
    ind=int(filename[:len(filename)-5])-36
    get_metrics(doc,ind)


# In[ ]:




